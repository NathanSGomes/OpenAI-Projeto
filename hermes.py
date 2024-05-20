import os
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import PyPDF2

load_dotenv()

# Configurações da OpenAI
client = OpenAI(
  api_key=os.environ.get('OPENAI_API_KEY'),  
)

# Pasta de origem e destino
pasta_texto = "0_Texto"
pasta_resumo = "1_Resumo"

# Função para ler o arquivo escolhido pelo usuário
def ler_arquivo(arquivo_escolhido):
    caminho_arquivo = os.path.join(pasta_texto, arquivo_escolhido)
    if arquivo_escolhido.endswith('.pdf'): # .pdf
        with open(caminho_arquivo, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            return ' '.join(page.extract_text() for page in pdf.pages)
        
    elif arquivo_escolhido.endswith('.docx'): # .docx
        doc = Document(caminho_arquivo)
        return ' '.join(paragraph.text for paragraph in doc.paragraphs)
    
    else:
        with open(caminho_arquivo, "r", encoding="utf-8") as file:
            return file.read()

# Função para gerar um resumo usando a OpenAI
def gerar_resumo(texto, max_tokens=4096):
    resumo = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
        {"role": "user", "content": f"Escreva um resumo e destaque as informações mais relevantes com o seguinte texto: {texto}"},
        ],
        temperature=0.3, # Define a escala Mais_Preciso (0) - MaisCriativo (1)
        max_tokens=max_tokens, # Define o tamanho do resumo desejado
        n=1, # Número de respostas a serem geradas
    )
    return resumo.choices[0].message.content

def main():
    # Lista os arquivos na pasta Texto e os ordena alfabeticamente
    arquivos_texto = sorted(os.listdir(pasta_texto))

    while True:
        # Exibe a lista de arquivos e permite que o usuário escolha
        print("\nArquivos disponíveis para resumo:")
        for i, arquivo in enumerate(arquivos_texto, start=1):
            print(f"{i}. {arquivo}")

        # Solicita a escolha do usuário
        escolha = input("Digite o número do arquivo que deseja resumir (ou 0 para sair): ")
        if escolha.isdigit():
            escolha = int(escolha)
            if escolha == 0:
                print("Programa encerrado.")
                break
            elif 1 <= escolha <= len(arquivos_texto):
                arquivo_escolhido = arquivos_texto[escolha - 1]
                print(f"Você escolheu o arquivo: {arquivo_escolhido}")
                confirmacao = input("Tem certeza que deseja gerar um resumo para este arquivo? (s/n): ")
                if confirmacao.lower() == "s":

                    # Lê o conteúdo do arquivo
                    conteudo = ler_arquivo(arquivo_escolhido)

                    # Gera o resumo
                    resumo = gerar_resumo(conteudo)

                    # Cria o diretório de destino se não existir
                    if not os.path.exists(pasta_resumo):
                        os.makedirs(pasta_resumo)

                    # Cria o arquivo de resumo na pasta Resumo
                    nome_resumo = f"Resumo_{os.path.splitext(arquivo_escolhido)[0]}.docx"
                    caminho_resumo = os.path.join(pasta_resumo, nome_resumo)
                    doc = Document()
                    doc.add_paragraph(resumo)
                    doc.save(caminho_resumo)
                    
                    print(f"\nResumo gerado com sucesso: {nome_resumo}")
                else:
                    print("\nOperação cancelada. Escolha outro arquivo.")
            else:
                print("\nNúmero inválido. Tente novamente.")
        else:
            print("\nEntrada inválida. Digite um número válido.")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print("\nOcorreu um erro:", e)